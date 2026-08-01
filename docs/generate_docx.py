"""Generate Phase 1 documentation as DOCX."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ── Title ──
title = doc.add_heading('DeepFake Detector — Phase 1: Preprocessing Pipeline', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ── 1. Architecture Overview ──
doc.add_heading('1. Architecture Overview', level=1)
doc.add_paragraph(
    'Phase 1 implements the foundational preprocessing pipeline that converts '
    'raw video files into aligned 299x299 face crops ready for Xception feature extraction. '
    'The pipeline consists of four sequential stages:'
)

stages = [
    ('Frame Extraction', 'Samples video at ~7.5 FPS, detects and discards blurry/corrupted frames, enhances moderately blurry ones.'),
    ('Face Detection', 'Finds the primary (largest, highest-confidence) face in each sampled frame using MediaPipe BlazeFace.'),
    ('Face Alignment', 'Runs MediaPipe FaceLandmarker (468 landmarks), rotates eyes to horizontal, crops with 25% margin, resizes to 299x299.'),
    ('Orchestration & Output', 'VideoPreprocessor ties all stages together. Saves face crops + metadata.json per video.'),
]
for name, desc in stages:
    p = doc.add_paragraph()
    runner = p.add_run(f'{name}: ')
    runner.bold = True
    p.add_run(desc)

# ── Data Flow Diagram ──
doc.add_heading('1.1 Data Flow', level=2)
flow_text = (
    'video.mp4\n'
    '    |\n'
    '    v\n'
    '+-----------------------+\n'
    '|  FrameExtractor       |  Sampled frames at ~7.5 FPS\n'
    '|  (frame_extractor.py) |  Blur detection & enhancement\n'
    '+-----------------------+\n'
    '    |\n'
    '    v\n'
    '+-----------------------+\n'
    '|  MediaPipeFaceDetector|  Primary face bounding box\n'
    '|  (face_detector.py)   |\n'
    '+-----------------------+\n'
    '    |\n'
    '    v\n'
    '+-----------------------+\n'
    '|  FaceAligner          |  468 landmarks, eye-level rotation,\n'
    '|  (face_alignment.py)  |  25% margin crop, 299x299 resize\n'
    '+-----------------------+\n'
    '    |\n'
    '    v\n'
    'aligned 299x299 face crops + metadata.json'
)
p = doc.add_paragraph()
run = p.add_run(flow_text)
run.font.name = 'Consolas'
run.font.size = Pt(8)

# ── Output Structure ──
doc.add_heading('1.2 Output Directory Structure', level=2)
out_text = (
    'data/processed/<video_name>/\n'
    '    frames/              (optional — extracted frames)\n'
    '        frame_0001.jpg\n'
    '        frame_0002.jpg\n'
    '        ...\n'
    '    faces/               (aligned 299x299 face crops)\n'
    '        face_0001.jpg\n'
    '        face_0002.jpg\n'
    '        ...\n'
    '    metadata.json        (video info, frame mapping, config)'
)
p = doc.add_paragraph()
run = p.add_run(out_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ── 2. File-by-File Documentation ──
doc.add_heading('2. File-by-File Documentation', level=1)

# ── Helper to add file section ──
def add_file_section(filename, path, purpose, imports, classes_functions, call_chain, notes=None):
    doc.add_heading(f'2.{file_counter[0]} {filename}', level=2)
    file_counter[0] += 1
    p = doc.add_paragraph()
    run = p.add_run('Path: ')
    run.bold = True
    p.add_run(f'src/preprocessing/{path}')

    p = doc.add_paragraph()
    run = p.add_run('Purpose: ')
    run.bold = True
    p.add_run(purpose)

    doc.add_heading('Imports', level=3)
    for imp in imports:
        p = doc.add_paragraph(imp, style='List Bullet')

    doc.add_heading('Classes & Functions', level=3)
    for cf in classes_functions:
        p = doc.add_paragraph()
        runner = p.add_run(cf['name'])
        runner.bold = True
        if cf.get('type'):
            p.add_run(f' — {cf["type"]}')
        if cf.get('desc'):
            doc.add_paragraph(cf['desc'])
        if cf.get('methods'):
            for m in cf['methods']:
                mp = doc.add_paragraph(m, style='List Bullet 2')

    if call_chain:
        doc.add_heading('Call Chain', level=3)
        for c in call_chain:
            doc.add_paragraph(c, style='List Bullet')

    if notes:
        doc.add_heading('Notes', level=3)
        for n in notes:
            doc.add_paragraph(n, style='List Bullet')

file_counter = [1]

# ── utils.py ──
add_file_section(
    'utils.py', 'utils.py',
    'Shared helpers used by all preprocessing modules — logging, filesystem, blur detection, frame enhancement, image resizing.',
    [
        'from __future__ import annotations',
        'import logging, os, urllib.request',
        'from pathlib import Path',
        'import cv2, numpy as np',
    ],
    [
        {'name': 'get_logger(name, level)', 'type': 'function', 'desc': 'Creates a module-level logger with a consistent format (timestamp, level, module name). Adds a StreamHandler only once to avoid duplicate messages.'},
        {'name': 'ensure_dir(path) -> Path', 'type': 'function', 'desc': 'Creates a directory (mkdir -p). Returns the Path.'},
        {'name': 'list_videos(directory, extensions) -> list[Path]', 'type': 'function', 'desc': 'Recursively finds all video files matching given extensions (mp4, avi, mov, mkv).'},
        {'name': 'ensure_model_asset(url, filename) -> Path', 'type': 'function', 'desc': 'Downloads a MediaPipe model from GCS to src/preprocessing/models/ on first run. Cached for subsequent runs.'},
        {'name': 'variance_of_laplacian(image) -> float', 'type': 'function', 'desc': 'Computes Laplacian variance — a standard focus measure. Higher = sharper.'},
        {'name': 'is_blurry(image, threshold=100.0) -> bool', 'type': 'function', 'desc': 'Returns True if variance_of_laplacian < threshold. Default 100.0 is a common starting point.'},
        {'name': 'is_severely_corrupted(image, threshold=15.0) -> bool', 'type': 'function', 'desc': 'Stricter blur check. Frames below this are dropped entirely (not enhanced).'},
        {'name': 'sharpen_frame(image, amount=1.5, radius=3) -> np.ndarray', 'type': 'function', 'desc': 'Unsharp-mask sharpening: Gaussian blur → subtract from original with weight.'},
        {'name': 'deblur_frame(image) -> np.ndarray', 'type': 'function', 'desc': 'Denoise (fastNlMeansDenoisingColored) + sharpen combo. Placeholder — can be swapped for learned model later.'},
        {'name': 'enhance_frame(image, blur_threshold=100.0) -> np.ndarray', 'type': 'function', 'desc': 'If blurry, deblur; otherwise return unchanged.'},
        {'name': 'resize_with_pad(image, target_size=299) -> np.ndarray', 'type': 'function', 'desc': 'Resize preserving aspect ratio, pad to square with black borders.'},
    ],
    [],
    [
        'PathLike = Union[str, os.PathLike] — central type alias',
        'Model cache: src/preprocessing/models/',
        'Designed to have zero inter-module dependencies (only OpenCV + NumPy)',
    ]
)

doc.add_page_break()

# ── face_detector.py ──
add_file_section(
    'face_detector.py', 'face_detector.py',
    'Face detection using the MediaPipe Tasks API (v1.0+). Detects the primary (largest, highest-confidence) face in each frame.',
    [
        'from __future__ import annotations',
        'from abc import ABC, abstractmethod',
        'from dataclasses import dataclass',
        'import numpy as np',
        'from .utils import ensure_model_asset, get_logger',
    ],
    [
        {'name': 'FaceBox', 'type': 'dataclass', 'desc': 'Bounding box with x, y, width, height, confidence. Has as_xyxy() and area() helpers.', 'methods': ['as_xyxy() -> (x1, y1, x2, y2)', 'area() -> width * height']},
        {'name': 'BaseFaceDetector', 'type': 'abstract class', 'desc': 'Interface so downstream code (FaceAligner, VideoPreprocessor) does not depend on a specific detector backend.', 'methods': ['detect(image) -> List[FaceBox] (abstract)', 'detect_primary(image) -> Optional[FaceBox] (picks max area × confidence)']},
        {'name': 'MediaPipeFaceDetector(BaseFaceDetector)', 'type': 'class', 'desc': 'Concrete implementation using MediaPipe Tasks FaceDetector. Downloads BlazeFace model on first use.', 'methods': ['__init__(min_confidence=0.5): downloads model, creates FaceDetectorOptions', 'detect(image): BGR → RGB → mp.Image → detect → parse detections', 'close() + context manager support']},
        {'name': 'get_default_detector(backend, **kwargs)', 'type': 'factory function', 'desc': 'Returns a detector instance by backend name. Currently only "mediapipe". Extensible for RetinaFace/MTCNN in later phases.'},
    ],
    [
        'MediaPipeFaceDetector.detect(image)',
        '    → mp.tasks.vision.FaceDetector.detect(mp_image)',
        '    → iterate results.detections → FaceBox per detection',
        '',
        'detect_primary() → max(faces, key=lambda f: f.area() * f.confidence)',
    ],
    [
        'Uses MediaPipe Tasks API (NOT the deprecated solutions API)',
        'Model auto-downloads from Google Cloud Storage on first run',
        'Model cached in src/preprocessing/models/blaze_face_short_range.tflite (~280 KB)',
        'Abstract BaseFaceDetector allows swapping backends without changing callers',
    ]
)

doc.add_page_break()

# ── face_alignment.py ──
add_file_section(
    'face_alignment.py', 'face_alignment.py',
    'Face alignment, cropping, and resizing. Given a frame + FaceBox, runs MediaPipe FaceLandmarker to get 468 landmarks, rotates eyes level, crops with margin, and resizes to 299x299.',
    [
        'from __future__ import annotations',
        'from dataclasses import dataclass',
        'import cv2, numpy as np',
        'from .face_detector import FaceBox',
        'from .utils import ensure_model_asset, get_logger, resize_with_pad',
    ],
    [
        {'name': 'AlignmentConfig', 'type': 'dataclass', 'desc': 'output_size=299, margin_fraction=0.25, min_landmark_confidence=0.5, min_tracking_confidence=0.5', 'methods': []},
        {'name': 'FaceAligner', 'type': 'class', 'desc': 'Aligns, crops, and resizes using MediaPipe FaceLandmarker (468 landmarks).', 'methods': [
            'align_and_crop(image, face_box) → Optional[np.ndarray]: main entry point; returns 299x299 BGR or None',
            '_get_landmarks(image, face_box): pads detection box by 50%, runs FaceLandmarker, returns (N,2) landmark array',
            '_rotate_to_level_eyes(image, landmarks): finds eye corners (indices 33, 263), computes angle, rotates image + landmarks via affine transform',
            '_crop_with_margin(image, landmarks): bounding box of landmarks + margin_fraction padding',
            'close() + context manager: releases FaceLandmarker resources',
        ]},
    ],
    [
        'FaceAligner.align_and_crop(frame, face_box)',
        '    → _get_landmarks() → (468, 2) array or None',
        '    → _rotate_to_level_eyes() → rotated image + transformed landmarks',
        '    → _crop_with_margin() → face region + 25% margin',
        '    → resize_with_pad() → 299x299 square',
    ],
    [
        'MediaPipe landmark indices: LEFT_EYE_OUTER=33, RIGHT_EYE_OUTER=263',
        'Rotation is about the midpoint between eyes',
        'Border mode: BORDER_REFLECT to avoid black corners after rotation',
        'Model: face_landmarker.task (~8 MB), auto-downloaded on first use',
    ]
)

doc.add_page_break()

# ── frame_extractor.py ──
add_file_section(
    'frame_extractor.py', 'frame_extractor.py',
    'Extracts frames from video at a target FPS, detects and handles blurry/corrupted frames, saves as JPEG.',
    [
        'from __future__ import annotations',
        'from dataclasses import dataclass, from pathlib import Path',
        'import cv2, numpy as np',
        'from .utils import PathLike, ensure_dir, enhance_frame, get_logger, is_severely_corrupted',
    ],
    [
        {'name': 'FrameExtractionConfig', 'type': 'dataclass', 'desc': 'target_fps=7.5, blur_threshold=100.0, corruption_threshold=15.0, enhance_blurry_frames=True, jpeg_quality=95', 'methods': []},
        {'name': 'FrameExtractor', 'type': 'class', 'desc': 'Opens video, iterates frames with sample interval, applies blur handling.', 'methods': [
            'extract(video_path, output_dir) → list[Path]: main entry; returns paths of saved frames',
            '_process_frame(frame): checks corruption → enhance if blurry → return or None',
            '_read_frames(cap): generator yielding raw frames',
        ]},
        {'name': 'extract_frames(video_path, output_dir, target_fps)', 'type': 'convenience function', 'desc': 'Wraps FrameExtractor + FrameExtractionConfig for simple usage.'},
    ],
    [
        'FrameExtractor.extract(video, output_dir)',
        '    → OpenCV VideoCapture → get source FPS',
        '    → sample_interval = round(source_fps / target_fps)',
        '    → for each sampled frame:',
        '        → is_severely_corrupted? drop',
        '        → is_blurry? enhance (sharpen + denoise)',
        '        → save as frame_NNNN.jpg',
    ],
    [
        'Default target FPS = 7.5 (midpoint of 5-10 FPS spec range)',
        'Moderately blurry frames are enhanced, not dropped',
        'Only severely corrupted frames (LoG variance < 15) are dropped',
    ]
)

doc.add_page_break()

# ── preprocess_video.py ──
add_file_section(
    'preprocess_video.py', 'preprocess_video.py',
    'Main pipeline orchestrator. Creates the three pipeline components once, loops over frames, and saves results with metadata.',
    [
        'from __future__ import annotations',
        'import json',
        'from dataclasses import dataclass, asdict',
        'from pathlib import Path',
        'import cv2, numpy as np',
        'from .face_alignment import FaceAligner, AlignmentConfig',
        'from .face_detector import MediaPipeFaceDetector, FaceBox',
        'from .frame_extractor import FrameExtractor, FrameExtractionConfig',
        'from .utils import PathLike, ensure_dir, get_logger, resize_with_pad',
    ],
    [
        {'name': 'PipelineConfig', 'type': 'dataclass', 'desc': 'Wraps FrameExtractionConfig + AlignmentConfig + flags (save_frames, save_faces, skip_frames_no_face). Uses __post_init__ to create default sub-configs if None.', 'methods': ['__post_init__(): creates default FrameExtractionConfig/AlignmentConfig if not provided']},
        {'name': 'ProcessedFrame', 'type': 'dataclass', 'desc': 'Stores paths to frame, face crop, bounding box, and whether a face was found.'},
        {'name': 'VideoResult', 'type': 'dataclass', 'desc': 'Aggregates all results for one video: paths, counts, list of ProcessedFrames.'},
        {'name': 'VideoPreprocessor', 'type': 'class (context manager)', 'desc': 'Creates FrameExtractor, MediaPipeFaceDetector, FaceAligner once, then processes frames in a loop.', 'methods': [
            'process(video_path, output_dir) → VideoResult: full pipeline',
            '_save_metadata(result, output_dir): writes metadata.json',
            'close() + __enter__/__exit__: releases MediaPipe resources',
        ]},
        {'name': 'process_video(video_path, output_dir, config)', 'type': 'convenience function', 'desc': 'Single-video entry point. Uses with-statement for automatic cleanup.'},
        {'name': 'process_directory(input_dir, output_dir, config, extensions)', 'type': 'convenience function', 'desc': 'Batch processing. Walks input_dir recursively, finds all videos by extension, processes each preserving directory structure.'},
    ],
    [
        'VideoPreprocessor.process(video, output_dir)',
        '    → FrameExtractor.extract() → list of frame paths',
        '    → for each frame:',
        '        → MediaPipeFaceDetector.detect_primary() → FaceBox or None',
        '        → if face: FaceAligner.align_and_crop() → 299x299',
        '        → save face crop, record in ProcessedFrame',
        '    → _save_metadata() → metadata.json',
        '',
        'process_video() → with VideoPreprocessor as p: p.process()',
        'process_directory() → for each video: process_video()',
    ],
    [
        'If save_frames=False, frames extracted to temp dir, then deleted',
        'If skip_frames_no_face=True (default), frames without faces produce no output',
        'If skip_frames_no_face=False, whole frame is resized to 299x299 as fallback',
        'metadata.json contains full config for reproducibility',
    ]
)

# ── __init__.py ──
doc.add_heading('2.6 src/preprocessing/__init__.py', level=2)
p = doc.add_paragraph()
run = p.add_run('Path: ')
run.bold = True
p.add_run('src/preprocessing/__init__.py')

p = doc.add_paragraph()
run = p.add_run('Purpose: ')
run.bold = True
p.add_run('Marks preprocessing as a Python package. Re-exports all public symbols for clean imports like from src.preprocessing import VideoPreprocessor.')

doc.add_paragraph('Exports:')
exports = [
    'face_detector: MediaPipeFaceDetector, FaceBox, get_default_detector',
    'face_alignment: FaceAligner, AlignmentConfig',
    'frame_extractor: FrameExtractor, FrameExtractionConfig, extract_frames',
    'preprocess_video: VideoPreprocessor, PipelineConfig, ProcessedFrame, VideoResult, process_video, process_directory',
    'utils: PathLike, ensure_dir, get_logger, list_videos, variance_of_laplacian, is_blurry, is_severely_corrupted, sharpen_frame, deblur_frame, enhance_frame, resize_with_pad',
]
for e in exports:
    doc.add_paragraph(e, style='List Bullet')

# ── Xception backbone ──
doc.add_heading('2.7 src/models/xception.py', level=2)
p = doc.add_paragraph()
run = p.add_run('Path: ')
run.bold = True
p.add_run('src/models/xception.py')

p = doc.add_paragraph()
run = p.add_run('Purpose: ')
run.bold = True
p.add_run('Xception backbone for feature extraction. Uses timm to load a pretrained Xception model with the classifier removed (num_classes=0).')

code_text = (
    'class XceptionBackbone(nn.Module):\n'
    '    def __init__(self):\n'
    '        self.model = timm.create_model("xception", pretrained=True, num_classes=0)\n'
    '    def forward(self, x):\n'
    '        return self.model(x)'
)
p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('Output shape: (batch, 2048) — the 2048-dim pooled features before the final classifier.')

doc.add_page_break()

# ── 3. Tests ──
doc.add_heading('3. Test Suite', level=1)
doc.add_paragraph(
    'The test suite (tests/test_pipeline.py) contains 11 tests covering unit-level logic '
    'and full integration scenarios. Run with: python -m pytest tests/'
)

doc.add_heading('3.1 Unit Tests (6)', level=2)
unit_tests = [
    ('test_pipeline_config_defaults', 'Verifies PipelineConfig creates default sub-configs when None is passed.'),
    ('test_pipeline_config_custom', 'Verifies custom config values override defaults correctly.'),
    ('test_processed_frame_dataclass', 'Verifies ProcessedFrame stores frame/face paths and FaceBox correctly.'),
    ('test_video_result_dataclass', 'Verifies VideoResult aggregates counts and paths.'),
    ('test_as_xyxy', 'Verifies FaceBox.as_xyxy() converts (x,y,w,h) → (x1,y1,x2,y2).'),
    ('test_area', 'Verifies FaceBox.area() returns width × height.'),
]
for name, desc in unit_tests:
    p = doc.add_paragraph()
    runner = p.add_run(name)
    runner.bold = True
    p.add_run(f': {desc}')

doc.add_heading('3.2 Integration Tests (4)', level=2)
int_tests = [
    ('test_process_video_creates_output', 'Creates a synthetic 10-frame video, runs the full pipeline, verifies output directories and metadata.json structure.'),
    ('test_process_video_no_frames_flag', 'Verifies that save_frames=False suppresses the frames/ directory.'),
    ('test_process_directory', 'Creates 2 synthetic videos, runs batch processing, verifies both are processed.'),
    ('test_pipeline_config_serialization', 'Verifies asdict(PipelineConfig) produces the expected nested dict.'),
]
for name, desc in int_tests:
    p = doc.add_paragraph()
    runner = p.add_run(name)
    runner.bold = True
    p.add_run(f': {desc}')

doc.add_heading('3.3 Error-Path Test (1)', level=2)
p = doc.add_paragraph()
runner = p.add_run('test_process_video_nonexistent')
runner.bold = True
p.add_run(': Verifies that process_video() raises FileNotFoundError for a nonexistent video path.')

doc.add_page_break()

# ── 4. Dependencies & Models ──
doc.add_heading('4. Dependencies & Models', level=1)

doc.add_heading('4.1 Python Dependencies', level=2)
deps = [
    ('OpenCV (opencv-python >= 4.10.0)', 'Video I/O, frame manipulation, image transforms, Laplacian blur detection'),
    ('MediaPipe (mediapipe >= 0.10.14 / 1.0.0+)', 'Face detection (BlazeFace) and face landmarking (FaceLandmarker) — Tasks API'),
    ('NumPy (numpy >= 2.1.0)', 'Array operations, landmark coordinate math'),
    ('PyTorch (torch >= 2.7.0)', 'Deep learning framework for Xception backbone'),
    ('timm (timm >= 1.0.19)', 'Pretrained Xception model loading'),
    ('Pillow', 'Image format support'),
    ('SciPy, pandas, scikit-learn', 'Utility libraries for future phases'),
    ('pytest', 'Test runner (dev dependency)'),
]
for name, desc in deps:
    p = doc.add_paragraph()
    runner = p.add_run(name)
    runner.bold = True
    p.add_run(f' — {desc}')

doc.add_heading('4.2 Downloaded Models', level=2)
p = doc.add_paragraph('Both models are downloaded on first use and cached at src/preprocessing/models/:')
models = [
    ('BlazeFace (short-range)', 'blaze_face_short_range.tflite', '~280 KB', 'Used by MediaPipeFaceDetector'),
    ('Face Landmarker', 'face_landmarker.task', '~8 MB', 'Used by FaceAligner for 468-landmark detection'),
]
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Shading Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Model'
hdr[1].text = 'Filename'
hdr[2].text = 'Size'
hdr[3].text = 'Used By'
for name, fname, size, used in models:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = fname
    row[2].text = size
    row[3].text = used

doc.add_page_break()

# ── 5. How to Run ──
doc.add_heading('5. How to Run', level=1)

doc.add_heading('5.1 Single Video', level=2)
p = doc.add_paragraph('python -m src.preprocessing.preprocess_video <video_path> --output-dir data/processed')
p = doc.add_paragraph('Example:')
p = doc.add_paragraph('python -m src.preprocessing.preprocess_video data/raw/fake/video.mp4 --output-dir data/processed --fps 10')

doc.add_heading('5.2 From Python', level=2)
code = (
    'from src.preprocessing import process_video, process_directory, PipelineConfig\n\n'
    '# Single video\n'
    'result = process_video("data/raw/fake/video.mp4", "data/processed")\n'
    'print(f"Faces: {result.num_faces_detected}")\n\n'
    '# Batch directory\n'
    'results = process_directory("data/raw", "data/processed")\n\n'
    '# Custom config\n'
    'config = PipelineConfig(save_frames=False)\n'
    'process_video("video.mp4", "out", config=config)'
)
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('5.3 Run Tests', level=2)
p = doc.add_paragraph('python -m pytest tests/ -v')
p = doc.add_paragraph()
run = p.add_run('Expected output: 11 passed')

# ── 6. MediaPipe API Migration Notes ──
doc.add_page_break()
doc.add_heading('6. MediaPipe API Migration (v0.x → v1.0+)', level=1)
doc.add_paragraph(
    'MediaPipe 1.0 replaced the legacy solutions API with the new Tasks API. '
    'The original codebase used the old API which no longer exists in modern MediaPipe installs. '
    'All changes are documented below for reference.'
)

doc.add_heading('6.1 Face Detection', level=2)
old_new_table = doc.add_table(rows=4, cols=2)
old_new_table.style = 'Light Shading Accent 1'
hdr = old_new_table.rows[0].cells
hdr[0].text = 'Old API (solutions)'
hdr[1].text = 'New API (tasks)'
data = [
    ('mp.solutions.face_detection.FaceDetection(...)', 'mp.tasks.vision.FaceDetector.create_from_options(...)'),
    ('detector.process(rgb)', 'detector.detect(mp.Image(...))'),
    ('det.score[0], det.location_data.relative_bounding_box', 'det.categories[0].score, det.bounding_box'),
]
for i, (old, new) in enumerate(data):
    row = old_new_table.rows[i + 1].cells
    row[0].text = old
    row[1].text = new

doc.add_heading('6.2 Face Landmarking', level=2)
old_new_table2 = doc.add_table(rows=5, cols=2)
old_new_table2.style = 'Light Shading Accent 1'
hdr2 = old_new_table2.rows[0].cells
hdr2[0].text = 'Old API (solutions)'
hdr2[1].text = 'New API (tasks)'
data2 = [
    ('mp.solutions.face_mesh.FaceMesh(...)', 'mp.tasks.vision.FaceLandmarker.create_from_options(...)'),
    ('self._mesh.process(rgb)', 'self._landmarker.detect(mp.Image(...))'),
    ('results.multi_face_landmarks[0].landmark', 'results.face_landmarks[0]'),
    ('min_detection_confidence', 'min_face_detection_confidence'),
]
for i, (old, new) in enumerate(data2):
    row = old_new_table2.rows[i + 1].cells
    row[0].text = old
    row[1].text = new

doc.add_page_break()

# ── 7. Project Structure ──
doc.add_heading('7. Phase 1 File Inventory', level=1)
file_table = doc.add_table(rows=1, cols=3)
file_table.style = 'Light Shading Accent 1'
hdr = file_table.rows[0].cells
hdr[0].text = 'File'
hdr[1].text = 'Lines'
hdr[2].text = 'Role'
files_info = [
    ('src/__init__.py', '0', 'Package marker'),
    ('src/models/__init__.py', '3', 'Package marker + XceptionBackbone export'),
    ('src/models/xception.py', '16', 'Xception backbone (feature extractor)'),
    ('src/preprocessing/__init__.py', '45', 'Package marker + re-exports'),
    ('src/preprocessing/utils.py', '212', 'Shared helpers (blur, logging, resize, model cache)'),
    ('src/preprocessing/face_detector.py', '139', 'Face detection (BlazeFace via MediaPipe Tasks)'),
    ('src/preprocessing/face_alignment.py', '183', 'Face alignment (FaceLandmarker + crop + resize)'),
    ('src/preprocessing/frame_extractor.py', '152', 'Frame extraction + blur handling'),
    ('src/preprocessing/preprocess_video.py', '221', 'Pipeline orchestrator'),
    ('tests/test_pipeline.py', '189', '11 tests (unit + integration + error-path)'),
]
for name, lines, role in files_info:
    row = file_table.add_row().cells
    row[0].text = name
    row[1].text = lines
    row[2].text = role

# ── Save ──
output_path = r'C:\Users\khark\OneDrive\Desktop\Projectbhai\deepfakedetector\docs\Phase1_Code_Documentation.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
